#!/usr/bin/env python3
"""
Regenerate analyses with updated prompt and create properly formatted Word document.
"""

import sys
import os
import logging
import re
from pathlib import Path
from typing import List, Tuple

# Add parent directory to path
sys.path.insert(0, str(Path(__file__).parent.parent))

from src.services.openai_service import OpenAIService
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Paths
DATA_DIR = Path(__file__).parent.parent / "cikkiro" / "data" / "MIK_Plenaris"
PROMPT_FILE = DATA_DIR / "prompt.txt"


def load_prompt() -> str:
    """Load analysis prompt from file."""
    if not PROMPT_FILE.exists():
        raise FileNotFoundError(f"Prompt file not found: {PROMPT_FILE}")

    prompt_text = PROMPT_FILE.read_text(encoding='utf-8')
    logger.info(f"Loaded updated prompt from {PROMPT_FILE}")
    return prompt_text


def get_transcription_files() -> List[Tuple[str, Path]]:
    """Get all transcription files and their titles."""
    # Get all transcribed.txt files (handles special dashes)
    files = sorted([f for f in DATA_DIR.glob("*.txt") if "transcribed.txt" in f.name])
    logger.info(f"Found {len(files)} transcription files")

    results = []
    for f in files:
        # Extract title from filename (remove everything after " - transcribed")
        name = f.name
        if " - transcribed.txt" in name:
            title = name.replace(" - transcribed.txt", "")
        else:
            continue  # Skip non-transcription files
        results.append((title, f))
        logger.info(f"  - {title}")

    return results


def parse_analysis_into_sections(analysis_text: str) -> List[Tuple[str, str]]:
    """Parse analysis text into sections (topic heading, content). Handles ##, ###, ####, etc."""
    sections = []
    current_topic = None
    current_content = []

    for line in analysis_text.split('\n'):
        # Match any markdown header level (##, ###, ####, etc.)
        header_match = re.match(r'^#+\s+(.+)$', line)
        if header_match:
            # Save previous section if exists
            if current_topic:
                sections.append((current_topic, '\n'.join(current_content).strip()))
            # Start new section
            current_topic = header_match.group(1).strip()
            current_content = []
        else:
            current_content.append(line)

    # Save last section
    if current_topic:
        sections.append((current_topic, '\n'.join(current_content).strip()))

    return sections


def process_bullet_content(content: str) -> List[str]:
    """Extract bullet points from content."""
    lines = content.split('\n')
    bullets = []

    for line in lines:
        line = line.strip()
        if line.startswith('- '):
            # Remove markdown formatting
            bullet_text = line[2:].strip()
            # Remove markdown bold/italic formatting
            bullet_text = bullet_text.replace('**', '')
            bullets.append(bullet_text)

    return bullets


def regenerate_analyses(openai_service: OpenAIService, prompt: str) -> List[Tuple[str, str]]:
    """Regenerate all analyses with updated prompt."""
    transcription_files = get_transcription_files()
    analyses = []

    for i, (title, transcription_path) in enumerate(transcription_files, 1):
        logger.info(f"\n[{i}/{len(transcription_files)}] Analyzing: {title}")

        try:
            # Read transcription
            transcript = transcription_path.read_text(encoding='utf-8')

            # Analyze with GPT-5.4
            logger.info("Running GPT-5.4 analysis...")
            analysis = openai_service.generate_chat_completion(
                transcript=transcript,
                user_prompt=prompt,
                model="gpt-5.4",
                max_tokens=10000
            )

            # Save analysis file
            analysis_path = transcription_path.parent / f"{title} - analysed.txt"
            analysis_path.write_text(analysis, encoding='utf-8')
            logger.info(f"✓ Analysis saved: {analysis_path.name}")

            analyses.append((title, analysis))

        except Exception as e:
            logger.error(f"✗ Failed to analyze {title}: {e}")
            continue

    logger.info(f"\n{'='*80}")
    logger.info(f"Regenerated {len(analyses)} analyses")
    logger.info(f"{'='*80}")

    return analyses


def add_heading_1(paragraph, text):
    """Add text as Heading 1 style."""
    paragraph.style = 'Heading 1'
    paragraph.clear()
    run = paragraph.add_run(text)
    return paragraph


def add_heading_2(paragraph, text):
    """Add text as Heading 2 style."""
    paragraph.style = 'Heading 2'
    paragraph.clear()
    run = paragraph.add_run(text)
    return paragraph


def create_formatted_word_document(analyses: List[Tuple[str, str]]) -> None:
    """Create Word document with proper formatting."""
    logger.info("\n" + "="*80)
    logger.info("CREATING FORMATTED WORD DOCUMENT")
    logger.info("="*80)

    doc = Document()

    # Add title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("MIK Plenáris Konferencia - AI Elemzések")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add metadata
    meta_para = doc.add_paragraph()
    meta_para.add_run("Készítés dátuma: ").bold = True
    meta_para.add_run("2026-04-09\n")
    meta_para.add_run("Elemzés módszere: ").bold = True
    meta_para.add_run("GPT-5.4 Chat Completion API\n")
    meta_para.add_run("Feldolgozott előadások: ").bold = True
    meta_para.add_run(f"{len(analyses)}\n")

    # Add each analysis with proper formatting
    for i, (title, analysis_text) in enumerate(analyses, 1):
        # Page break between presentations (except first)
        if i > 1:
            doc.add_page_break()

        # Add presentation title as Heading 1
        heading = doc.add_paragraph()
        add_heading_1(heading, f"{i}. {title}")

        # Parse analysis into sections
        sections = parse_analysis_into_sections(analysis_text)

        for topic, content in sections:
            # Add topic as Heading 2
            topic_para = doc.add_paragraph()
            add_heading_2(topic_para, topic)

            # Extract and add bullet points
            bullets = process_bullet_content(content)

            if bullets:
                for bullet_text in bullets:
                    # Add as bullet point
                    bullet_para = doc.add_paragraph(bullet_text, style='List Bullet')
            else:
                # If no bullets found, add content as regular paragraph
                if content.strip():
                    doc.add_paragraph(content)

    # Save document
    output_path = DATA_DIR / "output.docx"
    doc.save(str(output_path))
    logger.info(f"✓ Word document saved: {output_path.name}")
    logger.info(f"  Location: {output_path}")


def main():
    """Main entry point."""
    try:
        logger.info("Starting analysis regeneration with updated prompt...")
        logger.info(f"Data directory: {DATA_DIR}")

        # Initialize OpenAI service
        openai_service = OpenAIService()

        # Load updated prompt
        prompt = load_prompt()
        logger.info(f"Prompt loaded ({len(prompt)} characters)")

        # Regenerate analyses
        analyses = regenerate_analyses(openai_service, prompt)

        if not analyses:
            logger.error("No analyses regenerated. Cannot create Word document.")
            return

        # Create formatted Word document
        create_formatted_word_document(analyses)

        logger.info("\n" + "="*80)
        logger.info("SUCCESS: All tasks completed!")
        logger.info("="*80)
        logger.info(f"\nOutput files updated in: {DATA_DIR}")
        logger.info(f"  - {len(analyses)} analysis files (*-analysed.txt) - REGENERATED")
        logger.info(f"  - 1 Word document (output.docx) - REGENERATED with proper formatting")

    except Exception as e:
        logger.error(f"Fatal error: {e}", exc_info=True)
        sys.exit(1)


if __name__ == "__main__":
    main()
