#!/usr/bin/env python3
"""
Batch transcription and analysis script for MIK Plenaris audio files.

This script:
1. Transcribes all .m4a files in data/MIK_Plenaris directory using Whisper API
2. Analyzes each transcription using GPT-5.4 with the provided prompt
3. Saves transcriptions and analyses as separate text files
4. Creates a Word document with all analyses
"""

import sys
import os
import logging
from pathlib import Path
from typing import List, Tuple

# Add parent directory to path
sys.path.insert(0, str(Path(__file__).parent.parent))

from src.services.openai_service import OpenAIService
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Paths
DATA_DIR = Path(__file__).parent.parent / "cikkiro" / "data" / "MIK_Plenaris"
OUTPUT_DIR = DATA_DIR
PROMPT_FILE = DATA_DIR / "prompt.txt"


def load_prompt() -> str:
    """Load analysis prompt from file."""
    if not PROMPT_FILE.exists():
        raise FileNotFoundError(f"Prompt file not found: {PROMPT_FILE}")

    prompt_text = PROMPT_FILE.read_text(encoding='utf-8')
    logger.info(f"Loaded prompt from {PROMPT_FILE}")
    return prompt_text


def get_audio_files() -> List[Path]:
    """Get all .m4a files from MIK_Plenaris directory."""
    if not DATA_DIR.exists():
        raise FileNotFoundError(f"Data directory not found: {DATA_DIR}")

    files = sorted(DATA_DIR.glob("*.m4a"))
    logger.info(f"Found {len(files)} audio files in {DATA_DIR}")

    for f in files:
        logger.info(f"  - {f.name}")

    return files


def get_output_name(audio_file: Path, suffix: str) -> Path:
    """Generate output filename based on audio file."""
    # Remove .m4a extension and add suffix with .txt
    base_name = audio_file.stem  # Without extension
    output_name = f"{base_name} - {suffix}.txt"
    return OUTPUT_DIR / output_name


def transcribe_file(openai_service: OpenAIService, audio_file: Path) -> str:
    """Transcribe audio file with Hungarian language."""
    logger.info(f"\n{'='*80}")
    logger.info(f"TRANSCRIBING: {audio_file.name}")
    logger.info(f"{'='*80}")

    try:
        transcript = openai_service.transcribe_audio(
            str(audio_file),
            language="hu"  # Hungarian
        )
        logger.info(f"✓ Transcription completed ({len(transcript)} characters)")
        return transcript

    except Exception as e:
        logger.error(f"✗ Transcription failed: {e}")
        raise


def analyze_transcript(openai_service: OpenAIService, transcript: str, prompt: str) -> str:
    """Analyze transcript with GPT-5.4 using provided prompt."""
    logger.info(f"Analyzing transcript with GPT-5.4...")

    try:
        # Combine prompt with transcript
        full_prompt = f"{prompt}\n\n### Transzkript\n{transcript}"

        analysis = openai_service.generate_chat_completion(
            transcript=transcript,
            user_prompt=prompt,
            model="gpt-5.4",
            max_tokens=10000
        )
        logger.info(f"✓ Analysis completed ({len(analysis)} characters)")
        return analysis

    except Exception as e:
        logger.error(f"✗ Analysis failed: {e}")
        raise


def save_text_file(content: str, output_path: Path) -> None:
    """Save content to text file."""
    output_path.write_text(content, encoding='utf-8')
    logger.info(f"✓ Saved: {output_path.name}")


def batch_process() -> List[Tuple[str, str]]:
    """Process all audio files: transcribe and analyze."""

    # Initialize OpenAI service
    openai_service = OpenAIService()

    # Load prompt
    prompt = load_prompt()
    logger.info(f"Prompt loaded ({len(prompt)} characters)")

    # Get audio files
    audio_files = get_audio_files()

    if not audio_files:
        logger.error("No audio files found!")
        return []

    # Process each file
    results = []  # List of (audio_name, analysis) tuples

    for i, audio_file in enumerate(audio_files, 1):
        try:
            logger.info(f"\n[{i}/{len(audio_files)}] Processing: {audio_file.name}")

            # Transcribe
            transcript = transcribe_file(openai_service, audio_file)

            # Save transcription
            transcription_output = get_output_name(audio_file, "transcribed")
            save_text_file(transcript, transcription_output)

            # Analyze
            analysis = analyze_transcript(openai_service, transcript, prompt)

            # Save analysis
            analysis_output = get_output_name(audio_file, "analysed")
            save_text_file(analysis, analysis_output)

            # Store for Word document
            results.append((audio_file.stem, analysis))

            logger.info(f"✓ Completed: {audio_file.name}")

        except Exception as e:
            logger.error(f"✗ Failed to process {audio_file.name}: {e}")
            continue

    logger.info(f"\n{'='*80}")
    logger.info(f"Batch processing completed: {len(results)}/{len(audio_files)} files processed")
    logger.info(f"{'='*80}")

    return results


def create_word_document(analyses: List[Tuple[str, str]]) -> None:
    """Create Word document with all analyses."""
    logger.info("\n" + "="*80)
    logger.info("CREATING WORD DOCUMENT")
    logger.info("="*80)

    doc = Document()

    # Add title
    title = doc.add_heading("MIK Plenáris Konferencia - AI Elemzések", level=1)
    title_format = title.paragraph_format
    title_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add metadata
    metadata = doc.add_paragraph()
    metadata.add_run("Készítés dátuma: 2026-04-09\n").bold = True
    metadata.add_run("Elemzés módszere: GPT-5.4 Chat Completion API\n").italic = True
    metadata.add_run(f"Feldolgozott előadások: {len(analyses)}\n")

    # Add each analysis as a section
    for i, (title_text, analysis_text) in enumerate(analyses, 1):
        # Add heading for each presentation
        doc.add_page_break()
        heading = doc.add_heading(f"{i}. {title_text}", level=2)

        # Add analysis content
        doc.add_paragraph(analysis_text)

    # Save document
    output_path = OUTPUT_DIR / "output.docx"
    doc.save(str(output_path))
    logger.info(f"✓ Word document saved: {output_path.name}")
    logger.info(f"  Location: {output_path}")


def main():
    """Main entry point."""
    try:
        logger.info("Starting batch transcription and analysis...")
        logger.info(f"Data directory: {DATA_DIR}")
        logger.info(f"Output directory: {OUTPUT_DIR}")

        # Process all files
        analyses = batch_process()

        if not analyses:
            logger.error("No analyses completed. Cannot create Word document.")
            return

        # Create Word document
        create_word_document(analyses)

        logger.info("\n" + "="*80)
        logger.info("SUCCESS: All tasks completed!")
        logger.info("="*80)
        logger.info(f"\nOutput files created in: {OUTPUT_DIR}")
        logger.info(f"  - {len(analyses)} transcription files (*-transcribed.txt)")
        logger.info(f"  - {len(analyses)} analysis files (*-analysed.txt)")
        logger.info(f"  - 1 Word document (output.docx)")

    except Exception as e:
        logger.error(f"Fatal error: {e}", exc_info=True)
        sys.exit(1)


if __name__ == "__main__":
    main()
